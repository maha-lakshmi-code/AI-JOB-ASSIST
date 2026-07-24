"""
resume_parser.py — Extract text from PDF and DOCX resume files.

Uses:
  - pypdf (replaces deprecated PyPDF2) for PDF parsing
  - python-docx for DOCX parsing
"""

import re


def parse_resume(filepath: str, ext: str) -> str:
    """
    Extract plain text from a PDF or DOCX resume file.

    Args:
        filepath: Absolute or relative path to the uploaded file.
        ext:      File extension without dot — 'pdf' or 'docx'.

    Returns:
        Cleaned text string, or empty string on failure.
    """
    text = ""

    try:
        if ext == "pdf":
            text = _parse_pdf(filepath)
        elif ext == "docx":
            text = _parse_docx(filepath)
        else:
            print(f"[ResumeParser] Unsupported extension: {ext}")
            return ""
    except Exception as e:
        print(f"[ResumeParser] Unexpected error parsing {filepath}: {e}")
        return ""

    # Clean: collapse whitespace
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _parse_pdf(filepath: str) -> str:
    """Extract text from a PDF using pypdf."""
    text = ""
    try:
        # Primary: pypdf (modern, actively maintained replacement for PyPDF2)
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text

    except ImportError:
        pass  # fall through to pdfplumber

    try:
        # Fallback: pdfplumber (if installed)
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        return text

    except ImportError:
        print("[ResumeParser] Neither pypdf nor pdfplumber is installed.")
        return ""


def _parse_docx(filepath: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    text = ""
    try:
        from docx import Document
        doc = Document(filepath)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        return text

    except ImportError:
        print("[ResumeParser] python-docx is not installed.")
        return ""
